import os
import json
import re
import base64
from io import BytesIO
from typing import Tuple, Optional, Dict, Any, List

import streamlit as st
from PIL import Image
import fitz  # PyMuPDF
from groq import Groq
import openai  # For OpenRouter API compatibility
import streamlit.components.v1 as components

try:
    import docx
except ImportError:
    st.error("python-docx library not installed. Install it with 'pip install python-docx' to support DOCX files.")
    docx = None

# =============================
# API Keys and Models (Hardcoded as per instructions; in production, use secrets)
# =============================
GROQ_API_KEY = "grok_api_key"
OPENROUTER_API_KEY = "openrouter_api_key"

# Models
GRADING_MODEL = "openai/gpt-oss-20b"
VISION_MODEL = "qwen/qwen2.5-vl-72b-instruct:free"

# Clients
groq_client = Groq(api_key=GROQ_API_KEY)
openrouter_client = openai.OpenAI(base_url="https://openrouter.ai/api/v1", api_key=OPENROUTER_API_KEY)

# Default timeout
DEFAULT_TIMEOUT = 1200

# =============================
# Utilities
# =============================
def ensure_png_rgb(img: Image.Image) -> Image.Image:
    if img.mode not in ("RGB", "RGBA"):
        img = img.convert("RGB")
    if img.mode == "RGBA":
        bg = Image.new("RGB", img.size, (255, 255, 255))
        bg.paste(img, mask=img.split()[-1])
        img = bg
    return img

def image_to_base64(pil_image: Image.Image) -> str:
    pil_image = ensure_png_rgb(pil_image)
    buf = BytesIO()
    pil_image.save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode("utf-8")

def api_chat_completion(client, model: str, messages: List[Dict[str, Any]], max_tokens: int, response_format: Optional[Dict] = None, label: str = "API Call") -> Tuple[str, str, bool]:
    try:
        payload = {
            "model": model,
            "messages": messages,
            "max_tokens": max_tokens,
        }
        if response_format:
            payload["response_format"] = response_format
        response = client.chat.completions.create(**payload)
        content = response.choices[0].message.content.strip()
        raw = str(response)
        return content, raw, True
    except Exception as e:
        st.error(f"{label} error: {e}")
        return "", str(e), False

def try_parse_json(text: str) -> Tuple[Optional[dict], Optional[str], str]:
    if not text:
        return None, "Empty text", text
    candidates: List[str] = []
    s = text.strip()
    candidates.append(s)
    fenced = re.findall(r"```(?:json)?\s*(.*?)\s*```", s, flags=re.DOTALL | re.IGNORECASE)
    candidates.extend([blk.strip() for blk in fenced])
    if "{" in s and "}" in s:
        start, end = s.find("{"), s.rfind("}")
        if start < end:
            candidates.append(s[start:end + 1].strip())
    def clean_attempts(c: str) -> List[str]:
        out = [c]
        out.append(re.sub(r"\\(?![\"\\/bfnrtu])", r"\\\\", c))
        out.append(c.replace("\r", "").replace("\t", "\\t"))
        out.append(re.sub(r',\s*}', '}', c))
        out.append(re.sub(r',\s*\]', ']', c))
        out.append(re.sub(r'"\s*:', '":', c))
        out.append(re.sub(r'([^\x00-\x7F]+)', lambda m: m.group(1).encode('unicode_escape').decode(), c))
        out.append(c.replace("'", '"'))
        return out
    for cand in candidates:
        for trial in clean_attempts(cand):
            try:
                obj = json.loads(trial)
                if isinstance(obj, dict):
                    return obj, None, trial
            except json.JSONDecodeError as e:
                st.warning(f"JSON parse attempt failed: {e} for candidate: {trial[:100]}...")
                continue
    default_json = {"per_question": {}, "total": 0, "max_total": 0}
    return default_json, "All JSON parse attempts failed, returning default", text

# =============================
# PDF → PIL pages (via PyMuPDF)
# =============================
def pdf_bytes_to_pil_pages(pdf_bytes: bytes, dpi: int = 400) -> List[Image.Image]:
    try:
        doc = fitz.Document(stream=pdf_bytes, filetype="pdf")
        pages: List[Image.Image] = []
        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        for page in doc:
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            pages.append(img)
        doc.close()
        return pages
    except Exception as e:
        st.error(f"PDF processing error: {e}")
        return []

# =============================
# Qwen OCR – For Question Paper and Student Answers (Using OpenRouter)
# =============================
def extract_text_with_qwen(image: Image.Image, prompt: str) -> str:
    try:
        image_b64 = image_to_base64(image)
        messages = [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{image_b64}"}},
                ],
            }
        ]
        content, raw, ok = api_chat_completion(openrouter_client, VISION_MODEL, messages, 4096, label="OpenRouter Qwen OCR")
        if not ok:
            st.error(f"Raw OCR response: {raw}")
            st.error("Check if Qwen model is accessible or try increasing DPI.")
        return content.strip() if ok else ""
    except Exception as e:
        st.error(f"Qwen OCR processing error: {e}")
        return ""

def extract_student_details_with_qwen(image: Image.Image) -> Dict[str, str]:
    prompt = """
    Extract student details from the first page of this answer sheet image into a key-value pair dictionary. 
    Look for fields like Univ. Roll No., Name of the Student, Class/Semester, Name of the Paper, Code of paper, Total No. of Pages written by candidate, Sign. of the Student, Date of Exam.
    Output only the JSON object with keys in lowercase and values as extracted text, without additional commentary.
    """
    text = extract_text_with_qwen(image, prompt)
    obj, err, used = try_parse_json(text)
    if err or not isinstance(obj, dict):
        st.error(f"Failed to parse student details: {err or 'No valid JSON output.'}")
        return {}
    return obj

def extract_question_paper_with_qwen(image: Image.Image) -> str:
    prompt = """
    Transcribe the exact readable text from this question paper image, preserving table structure, questions, subparts, and OR options. Output as structured plain text, using markdown tables if applicable. Include only the question numbers (e.g., Q.1, Q.2) and question texts, ignoring headers like PRN, instructions, or metadata (Marks, CO, BL). Ensure OR options are clearly distinguished (e.g., label as Q1.a and Q1.b for OR alternatives). No commentary.
    """
    return extract_text_with_qwen(image, prompt)

def extract_student_answers_with_qwen(image: Image.Image) -> str:
    prompt = "Transcribe the exact readable text from this answer sheet image, starting from the second page onwards, ignoring the first page with student details. No commentary. Output plain text only."
    return extract_text_with_qwen(image, prompt)

def extract_text_from_file(uploaded_file, return_image: bool = False, is_question_paper: bool = False) -> Tuple[str, Optional[Image.Image], Optional[Dict[str, str]]]:
    if uploaded_file is None:
        return ("", None, None)
    mime = uploaded_file.type or ""
    text = ""
    preview_img: Optional[Image.Image] = None
    student_details: Optional[Dict[str, str]] = None
    try:
        uploaded_file.seek(0)  # Reset file pointer
        if mime == "application/pdf" or uploaded_file.name.lower().endswith(".pdf"):
            pdf_bytes = uploaded_file.read()
            pages = pdf_bytes_to_pil_pages(pdf_bytes, dpi=400)
            if not pages:
                st.error("No pages extracted from PDF. Please check the file.")
                return ("", None, None)
            ocr_texts = []
            if pages:
                student_details = extract_student_details_with_qwen(pages[0]) if not is_question_paper else {}
                for p in pages[1:] if not is_question_paper else pages:  # Start from second page for answers, all pages for question paper
                    if is_question_paper:
                        t = extract_question_paper_with_qwen(p)
                    else:
                        t = extract_student_answers_with_qwen(p)
                    if t:
                        ocr_texts.append(t)
            text = "\n".join(ocr_texts).strip()
            if return_image and pages:
                preview_img = pages[0] if not is_question_paper else pages[1]  # Preview first page for answers
        elif mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or uploaded_file.name.lower().endswith(".docx"):
            if docx is None:
                raise ImportError("docx library not available.")
            doc = docx.Document(uploaded_file)
            all_text = "\n".join([para.text for para in doc.paragraphs]).strip()
            if not is_question_paper:
                lines = all_text.split("\n")
                student_details_raw = "\n".join(lines[:5])  # Adjust based on typical layout
                student_details = extract_student_details_with_qwen(Image.fromarray(lines[:5].encode())) if student_details_raw else {}
                text = "\n".join(lines[5:]).strip()
            else:
                text = all_text
            preview_img = None
        elif mime == "text/plain" or uploaded_file.name.lower().endswith(".txt"):
            text = uploaded_file.read().decode("utf-8", errors="replace").strip()
            preview_img = None
        else:
            img = Image.open(uploaded_file)
            if is_question_paper:
                text = extract_question_paper_with_qwen(img).strip()
            else:
                student_details = extract_student_details_with_qwen(img)
                text = extract_student_answers_with_qwen(img).strip()
            if return_image:
                preview_img = ensure_png_rgb(img)
    except Exception as e:
        st.error(f"Error processing file: {e}")
    return (text, preview_img, student_details)

# =============================
# Improved Student Answer Parsing
# =============================
subpart_regex = re.compile(
    r"^\s*(?:(?:Ans|Answer|ans|answer)?\s*([a-z])[).:]\s*|(?:\d+\s*[(]([a-z])[)]\s*)|([a-z])\.\s*|\b([a-z])\)\s*)(.*)",
    re.IGNORECASE | re.MULTILINE
)

def parse_student_answers(raw_text: str) -> Dict[str, str]:
    answers: Dict[str, str] = {}
    current_q: Optional[str] = None
    current_key: Optional[str] = None
    current_answer: List[str] = []
    subpart_counter = 0
    subpart_letters = 'abcdefghijklmnopqrstuvwxyz'

    lines = raw_text.replace('\r\n', '\n').splitlines()
    skip_lines = 0
    for i, line in enumerate(lines):
        if skip_lines > 0:
            skip_lines -= 1
            continue
        line = line.strip()
        if not line:
            continue

        q_match = re.match(r"^(?:Q\.?\s*(\d+)|(\d+)\s*\.?)\s*[\):]?\s*(.*)", line, flags=re.IGNORECASE)
        if q_match:
            if current_key and current_answer:
                answers[current_key] = " ".join(current_answer).strip()
                current_answer = []
            current_q = f"Q{q_match.group(1) or q_match.group(2)}"
            current_key = None
            subpart_counter = 0
            remainder = q_match.group(3).strip() if q_match.group(3) else ""
            if remainder:
                subpart_match = subpart_regex.match(remainder)
                if subpart_match and current_q:
                    subq = next(filter(None, subpart_match.groups()[:4])).lower()
                    current_key = f"{current_q}.{subq}"
                    ans_text = subpart_match.group(5).strip() if subpart_match.group(5) is not None else ""
                    ans_text = re.sub(r"^(?:Define|What is|Name any|Distinguish between|Explain|Discuss|Answer|Ans|answer|ans)[:.]?\s*", "", ans_text, flags=re.IGNORECASE)
                    if ans_text:
                        current_answer.append(ans_text)
                    subpart_counter = max(subpart_counter, subpart_letters.index(subq) + 1)
                else:
                    if remainder and current_q:
                        current_key = f"{current_q}.{subpart_letters[subpart_counter]}"
                        ans_text = re.sub(r"^(?:Define|What is|Name any|Distinguish between|Explain|Discuss|Answer|Ans|answer|ans)[:.]?\s*", "", remainder, flags=re.IGNORECASE)
                        if ans_text:
                            current_answer.append(ans_text)
                        subpart_counter += 1
            continue

        subpart_match = subpart_regex.match(line)
        if subpart_match and current_q:
            if current_key and current_answer:
                answers[current_key] = " ".join(current_answer).strip()
                current_answer = []
            subq = next(filter(None, subpart_match.groups()[:4])).lower()
            current_key = f"{current_q}.{subq}"
            ans_text = subpart_match.group(5).strip() if subpart_match.group(5) is not None else ""
            ans_text = re.sub(r"^(?:Define|What is|Name any|Distinguish between|Explain|Discuss|Answer|Ans|answer|ans)[:.]?\s*", "", ans_text, flags=re.IGNORECASE)
            if ans_text:
                current_answer.append(ans_text)
            else:
                next_lines = []
                for j in range(i + 1, min(i + 4, len(lines))):
                    next_line = lines[j].strip()
                    if next_line and not re.match(r"^(?:Q\.?\s*\d+|Ans|Answer|ans|answer|\d+\s*\.?|[a-z][).:])", next_line, flags=re.IGNORECASE):
                        next_lines.append(next_line)
                    else:
                        break
                if next_lines:
                    current_answer.extend(next_lines)
                    skip_lines = len(next_lines)
            subpart_counter = max(subpart_counter, subpart_letters.index(subq) + 1)
            continue

        if current_q and line and not re.match(r"^(?:Define|What is|Name any|Distinguish between|Explain|Discuss)[:.]?\s*", line, flags=re.IGNORECASE):
            if not current_key:
                current_key = f"{current_q}.{subpart_letters[subpart_counter]}"
                subpart_counter += 1
            current_answer.append(line)

    if current_key and current_answer:
        answers[current_key] = " ".join(current_answer).strip()

    for k, v in list(answers.items()):
        if v:
            answers[k] = re.sub(r"\s+", " ", v).strip()
        else:
            answers[k] = ""

    return answers

# =============================
# Improved Grading with Conceptual Evaluation
# =============================
def evaluate_answer_quality(student_answer: str, reference: str, max_marks: int, subject: str, grade_level: str) -> Tuple[int, str]:
    if not student_answer.strip():
        return 0, "Not attempted."
    
    prompt = f"""
You are an expert examiner for {subject} (Grade {grade_level}).

### Instructions:
- Evaluate the student's answer using a chain-of-thought approach.
- Identify key concepts in the reference.
- Assess the student answer on:
  - **Accuracy** (0-10): Correctness of facts, definitions, and concepts.
  - **Depth** (0-10): Inclusion of key points, explanations, examples, and relevant details.
  - **Clarity and Structure** (0-10): Coherence, logical flow, and readability.
  - **Relevance** (0-10): Alignment with the reference question/answer, avoiding irrelevant information.
- Calculate a total score by averaging the criteria scores, then scale to max_marks ({max_marks}).
- Provide a detailed explanation, including sub-scores, quoting the student answer and reference, and justify the final score.
- Output JSON with "score" (final scaled integer), "sub_scores" (dict of criteria), and "explanation".

### Chain of Thought:
1. List key elements from reference.
2. Compare student answer point-by-point.
3. Assign sub-scores with reasoning.
4. Compute average and scale to max_marks.

### Student Answer:
{student_answer}

### Reference (Question or Sample Answer):
{reference}

### Output Format:
{{
  "score": int,
  "sub_scores": {{"accuracy": int, "depth": int, "clarity": int, "relevance": int}},
  "explanation": "Detailed evaluation including CoT..."
}}
"""
    messages = [{"role": "user", "content": prompt}]
    content, raw, ok = api_chat_completion(groq_client, GRADING_MODEL, messages, 1024, None, label="Groq Grading Evaluation")
    if not ok:
        st.error(f"Raw evaluation response: {raw}")
        student_words = set(student_answer.lower().split())
        reference_words = set(reference.lower().split())
        common_words = student_words.intersection(reference_words)
        match_ratio = len(common_words) / max(1, len(reference_words))
        score = min(int(match_ratio * max_marks + 0.5), max_marks)
        explanation = f"Fallback keyword match ratio: {match_ratio:.2f}. Common words: {', '.join(common_words) if common_words else 'None'}."
        return score, explanation
    
    obj, err, used = try_parse_json(content)
    if err or not isinstance(obj, dict):
        st.error(f"Evaluation parsing failed: {err}")
        student_words = set(student_answer.lower().split())
        reference_words = set(reference.lower().split())
        common_words = student_words.intersection(reference_words)
        match_ratio = len(common_words) / max(1, len(reference_words))
        score = min(int(match_ratio * max_marks + 0.5), max_marks)
        explanation = f"Fallback keyword match ratio: {match_ratio:.2f}. Common words: {', '.join(common_words) if common_words else 'None'}."
        return score, explanation
    
    score = min(max(int(obj.get("score", 0)), 0), max_marks)
    sub_scores = obj.get("sub_scores", {"accuracy": 0, "depth": 0, "clarity": 0, "relevance": 0})
    explanation = obj.get("explanation", "Assessed.").strip()
    return score, f"Sub-scores: {sub_scores}. {explanation}"

# =============================
# Improved Parsing of Question Paper (Using Groq)
# =============================
def parse_questions_with_llama(question_paper_text: str, subject: str, grade_level: str) -> Dict[str, str]:
    if not question_paper_text.strip():
        st.error("Question paper text is empty.")
        return {}
    prompt = f"""
You are an expert examiner for {subject} (Grade {grade_level}).

### Instructions:
- Analyze the provided question paper text, which may be in a table format with columns like Q. No., Question, Marks, CO, BL.
- Extract the question text for each question and subpart accurately, handling OR alternatives as subparts (e.g., Q1.a for first option, Q1.b for OR alternative).
- Ensure all questions and subparts are captured completely, even if the text is split across lines or tables.
- Ignore table artifacts like | --- |, PRN, instructions, headers (e.g., Q.No., Marks, CO, BL); focus only on question numbers and their corresponding texts.
- For questions with OR options, explicitly label them as Q1.a and Q1.b, ensuring no question is missed.
- If no OR, use "Q1": "Full question text".
- Validate that each question number (Q1, Q2, etc.) has at least one corresponding question text.
- Output ONLY the JSON object, without any additional text, commentary, or markdown.

### Question Paper Text:
{question_paper_text}

### Required Output Format:
{{
  "Q1.a": "Question text for first option",
  "Q1.b": "Question text for OR option",
  "Q2": "Full question text",
  ...
}}
"""
    messages = [{"role": "user", "content": prompt}]
    content, raw, ok = api_chat_completion(groq_client, GRADING_MODEL, messages, 4096, {"type": "json_object"}, label="Groq Question Parsing")
    if not ok:
        st.error(f"Raw parsing response: {raw}")
        return {}
    obj, err, used = try_parse_json(content)
    if err:
        st.error(f"Raw content before parsing: {content[:500]}...")
        return {}
    if not obj or not isinstance(obj, dict):
        st.error(f"Failed to parse questions: {err or 'No valid JSON output.'}")
        return {}
    # Validate all question numbers have text
    expected_questions = set(f"Q{i}" for i in range(1, 6))  # Adjust range based on expected questions
    parsed_questions = set(k.split('.')[0] for k in obj.keys())
    missing_questions = expected_questions - parsed_questions
    if missing_questions:
        st.warning(f"Missing questions detected: {missing_questions}. Please check the question paper text.")
    return obj

# =============================
# Parsing of Sample Paper (Using Groq)
# =============================
def parse_sample_answers_with_llama(sample_paper_text: str, subject: str, grade_level: str) -> Dict[str, str]:
    if not sample_paper_text.strip():
        st.error("Sample paper text is empty.")
        return {}
    prompt = f"""
You are an expert examiner for {subject} (Grade {grade_level}).

### Instructions:
- Analyze the provided sample paper text, which may include table structures or markdown-like formatting.
- Extract the correct answers for each question and subpart (e.g., Q1.a, Q1.b).
- For answers with OR alternatives, treat the first as subpart 'a' and the "OR" alternative as subpart 'b', using keys like "Q1.a" and "Q1.b".
- Ignore table headers like Q.No., Marks, CO, BL; focus on question numbers and answer texts.
- Do NOT extract marks; only focus on the answer text.
- Ensure answers are concise and match the sample paper exactly.
- Output ONLY the JSON object, without any additional text, commentary, or markdown.

### Sample Paper Text:
{sample_paper_text}

### Required Output Format:
{{
  "Q1.a": "Correct answer text",
  "Q1.b": "Correct answer text",
  ...
}}
"""
    messages = [{"role": "user", "content": prompt}]
    content, raw, ok = api_chat_completion(groq_client, GRADING_MODEL, messages, 4096, {"type": "json_object"}, label="Groq Sample Answer Parsing")
    if not ok:
        st.error(f"Raw parsing response: {raw}")
    obj, err, used = try_parse_json(content)
    if err:
        st.error(f"Raw content before parsing: {content}")
    if not obj or not isinstance(obj, dict):
        st.error(f"Failed to parse sample answers: {err or 'No valid JSON output.'}")
        return {}
    return obj

# =============================
# Sequential Grading per Question
# =============================
def grade_question(q: str, ms: Dict[str, Any], student_answers: Dict[str, str], sample_answers: Dict[str, str], question_texts: Dict[str, str], subject: str, grade_level: str) -> Dict[str, Dict[str, Any]]:
    subparts = [k for k, v in ms.items() if isinstance(v, dict) and "max" in v and re.fullmatch(r"[a-z]", str(k))]
    q_results: Dict[str, Dict[str, Any]] = {}
    for sp in subparts:
        key = f"{q}.{sp}"
        max_marks = int(ms.get(sp, {}).get("max", 0))
        student_answer = student_answers.get(key, "")
        reference = sample_answers.get(key, question_texts.get(key, ""))
        if not reference and key in question_texts:
            st.warning(f"No sample answer found for {key}. Using question text as reference.")
            reference = question_texts[key]
        elif not reference:
            st.error(f"No reference found for {key}. Skipping evaluation for this subpart.")
            continue
        score, explanation = evaluate_answer_quality(student_answer, reference, max_marks, subject, grade_level)
        q_results[sp] = {"score": score, "max": max_marks, "explanation": explanation, "reference": reference}
    return q_results

def apply_marking_rules(q: str, ms: Dict[str, Any], q_results: Dict[str, Dict[str, Any]]) -> Tuple[Dict[str, Dict[str, Any]], int, int]:
    subparts = list(q_results.keys())
    q_score = 0
    q_max = 0
    selected_subs = {}
    if ms.get("attempt_any_one", False) or ms.get("OR", True):  # Assume OR for this question paper
        if subparts:
            best_sp = max(subparts, key=lambda s: q_results[s]["score"])
            selected_subs = {best_sp: q_results[best_sp]}
            q_score += q_results[best_sp]["score"]
            q_max += max([q_results[s]["max"] for s in subparts])
    else:
        selected_subs = q_results
        for sp in subparts:
            q_score += q_results[sp]["score"]
            q_max += q_results[sp]["max"]
    return selected_subs, q_score, q_max

def grade_sequentially(student_text: str,
                       sample_answers: Dict[str, str],
                       question_texts: Dict[str, str],
                       marking_scheme: Dict[str, Any],
                       subject: str,
                       grade_level: str) -> Dict[str, Any]:
    student_answers = parse_student_answers(student_text)
    if not student_answers:
        st.error("No student answers parsed. Check OCR output in 'Debug: Raw OCR Student Text'.")
        return {"per_question": {}, "total": 0, "max_total": 0}

    per_q = {}
    total = 0
    max_total = 0
    questions = sorted([q for q in marking_scheme if q != "grand_total" and isinstance(marking_scheme[q], dict)], key=lambda x: int(x[1:]))
    
    for q in questions:
        ms = marking_scheme[q]
        q_results = grade_question(q, ms, student_answers, sample_answers, question_texts, subject, grade_level)
        if q_results:
            selected_subs, q_score, q_max = apply_marking_rules(q, ms, q_results)
            per_q[q] = selected_subs
            total += q_score
            max_total += q_max

    return {
        "per_question": per_q,
        "total": total,
        "max_total": max_total,
        "expected_max_total": marking_scheme.get("grand_total", max_total)
    }

# =============================
# Streamlit UI
# =============================
st.set_page_config(page_title="OCR & Grading", layout="wide")
st.title("📘 Student Answer OCR & Grading")

with st.sidebar:
    st.header("API Configurations")
    st.caption(f"Grading Model: {GRADING_MODEL} (Groq)")
    st.caption(f"Vision OCR Model: {VISION_MODEL} (OpenRouter)")
    st.warning("API keys are hardcoded for this demo. In production, use Streamlit secrets.")
    st.warning("Ensure no local 'groq.py' file is conflicting with the groq package.")

# Initialize session state
if "question_text" not in st.session_state:
    st.session_state.question_text = ""
if "sample_text" not in st.session_state:
    st.session_state.sample_text = ""
if "marking_scheme" not in st.session_state:
    st.session_state.marking_scheme = {}
if "question_texts" not in st.session_state:
    st.session_state.question_texts = {}
if "sample_answers" not in st.session_state:
    st.session_state.sample_answers = {}
if "teacher_artifacts_uploaded" not in st.session_state:
    st.session_state.teacher_artifacts_uploaded = False
if "pages" not in st.session_state:
    st.session_state.pages = []
if "accumulated_student_text" not in st.session_state:
    st.session_state.accumulated_student_text = ""
if "student_details" not in st.session_state:
    st.session_state.student_details = {}
if "last_interaction" not in st.session_state:
    st.session_state.last_interaction = 0

# Step 1: Upload Teacher Artifacts
st.header("Step 1: Upload Question Paper or Sample Answer Sheet and Marking Scheme (One-Time)")
if not st.session_state.teacher_artifacts_uploaded:
    uploaded_question_paper = st.file_uploader("Upload Question Paper (DOCX/PDF/TXT) - Optional", type=["docx", "pdf", "txt"], key="qp")
    uploaded_answer_sheet = st.file_uploader("Upload Teacher Answer Sheet (TXT/PDF/IMG) - Optional", type=["txt", "pdf", "png", "jpg", "jpeg"], key="ans")
    uploaded_marking_scheme = st.file_uploader("Upload Marking Scheme (JSON)", type=["json"], key="ms")

    if st.button("Submit Teacher Artifacts"):
        if uploaded_marking_scheme:
            try:
                uploaded_marking_scheme.seek(0)
                st.session_state.marking_scheme = json.load(uploaded_marking_scheme)
            except Exception as e:
                st.error(f"Failed to load marking scheme: {e}")
        if uploaded_answer_sheet:
            try:
                st.session_state.sample_text, _, _ = extract_text_from_file(uploaded_answer_sheet, is_question_paper=False)
                if st.session_state.sample_text:
                    st.session_state.sample_answers = parse_sample_answers_with_llama(
                        st.session_state.sample_text, subject="Fundamentals of Management and Strategies Formulation", grade_level="T.Y B.Tech. SEM-V"
                    )
                else:
                    st.error("No text extracted from teacher answer sheet.")
            except Exception as e:
                st.error(f"Failed to process teacher answer sheet: {e}")
        if uploaded_question_paper:
            try:
                st.session_state.question_text, _, _ = extract_text_from_file(uploaded_question_paper, is_question_paper=True)
                if st.session_state.question_text:
                    st.session_state.question_texts = parse_questions_with_llama(
                        st.session_state.question_text, subject="Fundamentals of Management and Strategies Formulation", grade_level="T.Y B.Tech. SEM-V"
                    )
                else:
                    st.error("No text extracted from question paper.")
            except Exception as e:
                st.error(f"Failed to process question paper: {e}")
        if st.session_state.marking_scheme:
            st.session_state.teacher_artifacts_uploaded = True
            st.success("Teacher artifacts uploaded successfully. Proceed to Step 2.")
        else:
            st.error("Marking scheme is required. Please upload a valid JSON file.")
else:
    st.info("Teacher artifacts already uploaded. Use the button below to reset if needed.")
    if st.button("Reset Teacher Artifacts"):
        st.session_state.teacher_artifacts_uploaded = False
        st.session_state.question_text = ""
        st.session_state.sample_text = ""
        st.session_state.marking_scheme = {}
        st.session_state.question_texts = {}
        st.session_state.sample_answers = {}
        st.session_state.accumulated_student_text = ""
        st.session_state.student_details = {}
        st.rerun()

with st.expander("Preview: Teacher Artifacts"):
    st.text_area("Teacher Answer Sheet (preview)", st.session_state.sample_text, height=180, disabled=True)
    st.subheader("Parsed Sample Answers")
    st.json(st.session_state.sample_answers)
    st.text_area("Question Paper (preview)", st.session_state.question_text, height=180, disabled=True)
    st.subheader("Parsed Questions")
    st.json(st.session_state.question_texts)
    st.subheader("Marking Scheme")
    st.json(st.session_state.marking_scheme)

# Step 2: Upload Student Answer Sheet
st.header("Step 2: Upload Student Answer Sheet")
uploaded_student_auto = st.file_uploader("Upload Student Answer Sheet", type=["pdf", "png", "jpg", "jpeg", "docx", "txt"], key="stu")
manual_student_text = st.text_area("Or paste Student Answers manually (optional)", height=180, on_change=lambda: st.session_state.update({"last_interaction": st.session_state.get("last_interaction", 0) + 1}))

# Initialize student_text
student_text = st.session_state.accumulated_student_text

# Button to process the student answer sheet
if st.button("Process Student Answer Sheet", disabled=not (uploaded_student_auto or manual_student_text.strip())):
    try:
        st.session_state.pop("current_page", None)
        st.session_state.pop("extracted_texts", None)
        st.session_state.pop("pages", None)
        st.session_state.accumulated_student_text = ""
        st.session_state.last_interaction = 0
        st.session_state.student_details = {}

        if manual_student_text.strip():
            st.session_state.accumulated_student_text = manual_student_text.strip()
            student_text = st.session_state.accumulated_student_text
            st.success("Manual student answers processed successfully.")
        else:
            mime = uploaded_student_auto.type or ""
            if mime == "application/pdf" or uploaded_student_auto.name.lower().endswith(".pdf"):
                uploaded_student_auto.seek(0)
                pdf_bytes = uploaded_student_auto.read()
                st.session_state.pages = pdf_bytes_to_pil_pages(pdf_bytes, dpi=400)
                if not st.session_state.pages:
                    st.error("No pages loaded from PDF. Please check the file.")
                else:
                    st.session_state.extracted_texts = [""] * len(st.session_state.pages)
                    st.session_state.current_page = 0
                    st.session_state.student_details = extract_student_details_with_qwen(st.session_state.pages[0]) if st.session_state.pages else {}
            else:
                student_text, student_image, student_details = extract_text_from_file(uploaded_student_auto, return_image=True, is_question_paper=False)
                st.session_state.accumulated_student_text = student_text
                st.session_state.student_details = student_details or {}
                if student_text:
                    st.success("Student answer sheet processed successfully.")
                else:
                    st.error("No text extracted from student answer sheet.")
                if student_image:
                    st.image(student_image, caption="Student Answer Sheet (Preview)")
        st.rerun()
    except Exception as e:
        st.error(f"Error processing student answer sheet: {e}")

# Handle PDF page by page
if "pages" in st.session_state and st.session_state.pages:
    current_page = st.session_state.get("current_page", 0)
    extracted_texts = st.session_state.get("extracted_texts", [""] * len(st.session_state.pages))

    st.subheader(f"Page {current_page + 1} of {len(st.session_state.pages)}")

    if 0 <= current_page < len(st.session_state.pages):
        st.image(st.session_state.pages[current_page], caption=f"Page {current_page + 1}")
    else:
        st.error(f"Invalid page index: {current_page}. Total pages: {len(st.session_state.pages)}")

    if extracted_texts[current_page] == "":
        with st.spinner("Extracting text from this page..."):
            if current_page == 0:
                st.session_state.student_details = extract_student_details_with_qwen(st.session_state.pages[current_page])
            else:
                extracted = extract_student_answers_with_qwen(st.session_state.pages[current_page])
                extracted_texts[current_page] = extracted
            st.session_state.extracted_texts = extracted_texts

    edited_text = st.text_area(
        "Extracted/Edited Text for this Page", 
        extracted_texts[current_page], 
        height=300,
        on_change=lambda: st.session_state.update({"last_interaction": st.session_state.get("last_interaction", 0) + 1})
    )
    extracted_texts[current_page] = edited_text
    st.session_state.extracted_texts = extracted_texts

    accumulated = "\n\n".join([t.strip() for t in extracted_texts[1:] if t.strip()])  # Exclude first page
    st.session_state.accumulated_student_text = accumulated
    student_text = accumulated

    if accumulated:
        st.subheader("Preview of Accumulated Extracted Text So Far")
        st.text_area("Accumulated Text (Continued from Previous Pages)", accumulated, height=300, disabled=True)

    if st.session_state.student_details:
        st.subheader("Student Details")
        st.json(st.session_state.student_details)
        st.download_button("Download Student Details (JSON)", json.dumps(st.session_state.student_details, indent=2, ensure_ascii=False),
                          file_name="student_details.json")

    if accumulated:
        with st.expander("Preview: Parsed Student Answer Mapping (from accumulated text)"):
            parsed_answers = parse_student_answers(accumulated)
            st.json(parsed_answers)

    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("Previous", disabled=current_page <= 0, key="prev_button", on_click=lambda: st.session_state.update({"last_interaction": st.session_state.get("last_interaction", 0) + 1})):
            st.session_state.current_page = max(0, current_page - 1)
            st.rerun()
    with col2:
        next_disabled = current_page >= len(st.session_state.pages) - 1
        if st.button("Next", disabled=next_disabled, key="next_button", on_click=lambda: st.session_state.update({"last_interaction": st.session_state.get("last_interaction", 0) + 1})):
            if current_page < len(st.session_state.pages) - 1:
                st.session_state.current_page = current_page + 1
                st.rerun()
    with col3:
        if st.button("Finish Processing", key="finish_button", on_click=lambda: st.session_state.update({"last_interaction": st.session_state.get("last_interaction", 0) + 1})):
            student_text = st.session_state.accumulated_student_text
            st.session_state.pop("current_page", None)
            st.session_state.pop("extracted_texts", None)
            st.session_state.pop("pages", None)
            st.session_state.last_interaction = 0
            st.rerun()

    # Auto-advance to next page after 3 seconds of inactivity
    components.html(f"""
    <script>
    let timeoutId;
    let interactionCount = {st.session_state.get("last_interaction", 0)};
    
    function resetTimer() {{
        clearTimeout(timeoutId);
        timeoutId = setTimeout(function() {{
            if (interactionCount === {st.session_state.get("last_interaction", 0)}) {{
                var buttons = window.parent.document.querySelectorAll('button');
                for (var i = 0; i < buttons.length; i++) {{
                    if (buttons[i].innerText === 'Next' && !buttons[i].disabled) {{
                        buttons[i].click();
                        break;
                    }}
                }}
            }}
        }}, 3000);
    }}
    
    document.addEventListener('mousemove', resetTimer);
    document.addEventListener('keypress', resetTimer);
    document.addEventListener('click', resetTimer);
    document.addEventListener('scroll', resetTimer);
    resetTimer();
    </script>
    """, height=0)

# Display student answers used for grading
if student_text:
    st.subheader("Final Student Answers (Used for Grading)")
    with st.expander("Debug: Raw OCR Student Text"):
        st.text(student_text)
    with st.expander("Preview: Parsed Student Answer Mapping"):
        parsed_answers = parse_student_answers(student_text)
        st.json(parsed_answers)
        if not parsed_answers:
            st.warning("No answers parsed. Check OCR output or answer sheet format.")

# Step 3: Grading
st.header("Step 3: Grade Student Answers")
can_grade = bool(student_text) and bool(st.session_state.marking_scheme)
if st.button("🚀 Grade Now", disabled=not can_grade, on_click=lambda: st.session_state.update({"last_interaction": st.session_state.get("last_interaction", 0) + 1})):
    if can_grade:
        with st.spinner("Grading in progress..."):
            result = grade_sequentially(
                student_text,
                st.session_state.sample_answers,
                st.session_state.question_texts,
                st.session_state.marking_scheme,
                subject="Fundamentals of Management and Strategies Formulation",
                grade_level="T.Y B.Tech. SEM-V"
            )
        st.subheader("Grading Report")
        
        per_q = result.get("per_question", {})
        total = result.get("total", 0)
        max_total = result.get("max_total", 0)
        expected_max_total = result.get("expected_max_total", max_total)
        
        student_answers = parse_student_answers(student_text)
        
        table_data = []
        
        questions = sorted(per_q.keys(), key=lambda x: int(x[1:]))
        for q in questions:
            subs = per_q[q]
            st.subheader(f"Question {q}")
            for sp, details in sorted(subs.items()):
                key = f"{q}.{sp}"
                question_text = st.session_state.question_texts.get(key, details.get("reference", "Question text not available"))
                student_ans = student_answers.get(key, "Not attempted")
                explanation = details.get("explanation", "No explanation provided")
                score = details.get("score", 0)
                max_marks = details.get("max", 0)
                
                st.write(f"{q} {sp}: \"{question_text}\"")
                st.write(f"Student answer: \"{student_ans}\"")
                st.write(f"Evaluated description: \"{explanation}\"")
                st.write(f"Grading and marks: {score} / {max_marks}")
                
                table_data.append({
                    "Question": f"{q}.{sp}",
                    "Question Text": question_text,
                    "Answer": student_ans,
                    "Grade": f"{score} / {max_marks}",
                    "Remark": explanation
                })
        
        st.subheader("Overall Score")
        st.write(f"Total: {total} / {max_total}")
        if expected_max_total != max_total:
            st.write(f"Expected Max Total: {expected_max_total}")
        
        st.subheader("Summary Table")
        st.dataframe(table_data)
        
        st.download_button("⬇️ Download Grading Report (JSON)", json.dumps(result, indent=2, ensure_ascii=False),
                          file_name="grading_report.json")
        
        # Display Student Details at the end
        if st.session_state.student_details:
            st.subheader("Student Details")
            st.json(st.session_state.student_details)
            st.download_button("Download Student Details (JSON)", json.dumps(st.session_state.student_details, indent=2, ensure_ascii=False),
                              file_name="student_details.json")
    else:
        st.warning("Please ensure Marking Scheme (from Step 1) and Student Answers (from Step 2) are provided. Question Paper and Teacher Answer Sheet are optional.")
else:
    if not can_grade:
        st.warning("Please ensure Marking Scheme (from Step 1) and Student Answers (from Step 2) are provided. Question Paper and Teacher Answer Sheet are optional.")